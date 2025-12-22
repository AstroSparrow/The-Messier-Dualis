from docx import Document
import os

FOLDER_PATH = "."  # folder with .py files
OUTPUT_FILE = "CS_Project_Code.docx"

doc = Document()
doc.add_heading("The Messier Dualis – Source Code", level=1)

for file in os.listdir(FOLDER_PATH):
    if file.endswith(".py"):
        doc.add_heading(file, level=2)
        doc.add_paragraph("-" * 40)

        with open(file, "r", encoding="utf-8") as f:
            code = f.read()
            p = doc.add_paragraph(code)
            p.style.font.name = "Consolas"

doc.save(OUTPUT_FILE)
print("Word file created successfully!")
exit()